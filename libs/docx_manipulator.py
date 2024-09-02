from docx import Document


class DocxGenerator():
    def __init__(self):
        self.document = Document()

    def add_title(self, text, level):
        self.document.add_heading(text, level)

    def add_paragraph(self, text):
        self.document.add_paragraph(text)

    def add_image(self, image_path):
        self.document.add_picture(image_path)

    def add_table_concentracao_sem_correcao(self, records: tuple):
        self.document.add_heading('Concentracao sem correção', 1)
        table = self.document.add_table(rows=1, cols=5)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'N°'
        hdr_cells[1].text = 'Massa padrão'
        hdr_cells[2].text = 'Fator de equivalência'
        hdr_cells[3].text = 'Volume de solução'
        hdr_cells[4].text = 'Concentração Final'

        for id, massa, fator, volume, concentracao in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(id)
            row_cells[1].text = str(massa) + ' mg'
            row_cells[2].text = str(fator) + ' g/mol'
            row_cells[3].text = str(volume) + ' mol/L'
            row_cells[4].text = str(concentracao)
    
    def add_table_concentracao_com_correcao(self, records: tuple):
        self.document.add_heading('Concentracao com correção', 1)
        table = self.document.add_table(rows=1, cols=6)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'N°'
        hdr_cells[1].text = 'Massa padrão'
        hdr_cells[2].text = 'Fator de correção'
        hdr_cells[3].text = 'Fator de equivalência'
        hdr_cells[4].text = 'Volume NAOH'
        hdr_cells[5].text = 'Concentração Final'

        for id, massa, correcao, equivalencia, volume, concentracao in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(id)
            row_cells[1].text = str(massa) + ' mg'
            row_cells[2].text = str(correcao) + ' g/mol'
            row_cells[3].text = str(equivalencia) + ' g/mol'
            row_cells[4].text = str(volume) + ' mol/L'
            row_cells[5].text = str(concentracao)
    
    def add_table_concentracao_molar(self, records: tuple):
        self.document.add_heading('Concentracao molar', 1)
        table = self.document.add_table(rows=1, cols=5)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'N°'
        hdr_cells[1].text = 'Massa soluto'
        hdr_cells[2].text = 'Massa molar'
        hdr_cells[3].text = 'Volume solução'
        hdr_cells[4].text = 'Concentração Final'

        for id, soluto, molar, volume, concentracao in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(id)
            row_cells[1].text = str(soluto) + ' mg'
            row_cells[2].text = str(molar) + ' g/mol'
            row_cells[3].text = str(volume) + ' mol/L'
            row_cells[4].text = str(concentracao)

    
    def save_doc(self, filename):
        self.document.save(f'static/docs/{filename}.docx')


def generate_random_docx():
    records1 = (
            (1, 153, 15, 4848, 125),
            (2, 100, 1, 500, 105),
        )
    
    records2 = (
            (1, 153, 15, 4848, 125, 100),
            (2, 100, 1, 500, 105, 90),
        )
    
    doc = DocxGenerator()
    doc.add_title('Documento de exemplo', 0)
    doc.add_image('static\images\imagem_exemplo.png')
    doc.add_table_concentracao_sem_correcao(records1)
    doc.add_table_concentracao_com_correcao(records2)
    doc.add_table_concentracao_molar(records1)
    doc.save_doc('teste01')